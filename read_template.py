# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

"""
读取并分析毕业设计模板文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def read_template(template_path):
    """读取Word模板并提取内容结构"""
    doc = Document(template_path)
    
    print("=" * 80)
    print("[模板分析] 毕业设计模板分析报告")
    print("=" * 80)
    print(f"\n文件路径: {template_path}")
    print(f"\n{'='*80}")
    
    # 1. 文档基本属性
    print("\n【一、文档基本信息】")
    print(f"  段落数量: {len(doc.paragraphs)}")
    print(f"  表格数量: {len(doc.tables)}")
    print(f"  节数量: {len(doc.sections)}")
    
    # 2. 页面设置（从第一个section获取）
    if doc.sections:
        section = doc.sections[0]
        print(f"\n【二、页面设置】")
        print(f"  页面宽度: {section.page_width.cm:.2f} cm")
        print(f"  页面高度: {section.page_height.cm:.2f} cm")
        print(f"  上边距: {section.top_margin.cm:.2f} cm")
        print(f"  下边距: {section.bottom_margin.cm:.2f} cm")
        print(f"  左边距: {section.left_margin.cm:.2f} cm")
        print(f"  右边距: {section.right_margin.cm:.2f} cm")
        print(f"  装订线: {section.gutter.cm:.2f} cm" if section.gutter else "  装订线: 无")
    
    # 3. 提取所有段落文本和样式
    print(f"\n{'='*80}")
    print("【三、文档内容结构】")
    print(f"{'='*80}\n")
    
    content_structure = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else "无样式"
            
            # 获取字体信息
            font_info = ""
            if para.runs:
                run = para.runs[0]
                font_size = run.font.size.pt if run.font.size else "未设置"
                font_name = run.font.name if run.font.name else "未设置"
                bold = run.font.bold
                font_info = f"(字体:{font_name}, 字号:{font_size}, 粗体:{bold})"
            
            # 判断是否为标题
            is_heading = 'Heading' in style_name or '标题' in style_name or 'head' in style_name.lower()
            
            prefix = "【标题】" if is_heading else "【正文】"
            indent = "  " * (int(style_name.replace('Heading ', '').replace('Title', '0')) if is_heading and style_name.replace('Heading ', '').replace('Title', '0').isdigit() else 0)
            
            line_info = f"{prefix}{indent} {style_name}: {text[:100]}{'...' if len(text) > 100 else ''} {font_info}"
            print(line_info)
            
            content_structure.append({
                'index': i,
                'type': 'heading' if is_heading else 'text',
                'style': style_name,
                'text': text,
                'font_info': font_info
            })
    
    # 4. 分析表格结构
    if doc.tables:
        print(f"\n{'='*80}")
        print("【四、表格信息】")
        print(f"{'='*80}\n")
        
        for t_idx, table in enumerate(doc.tables):
            print(f"表格 {t_idx + 1}: {len(table.rows)} 行 × {len(table.columns)} 列")
            
            # 显示表头（第一行）
            if table.rows:
                header_cells = [cell.text.strip()[:20] for cell in table.rows[0].cells]
                print(f"  表头: {' | '.join(header_cells)}")
                
                # 显示前几行数据
                for row_idx, row in enumerate(table.rows[1:4], 1):  # 只显示前3行数据
                    cells = [cell.text.strip()[:15] for cell in row.cells]
                    print(f"  第{row_idx}行: {' | '.join(cells)}")
                if len(table.rows) > 4:
                    print(f"  ... 共{len(table.rows)}行数据")
            print()
    
    return {
        'doc': doc,
        'paragraph_count': len(doc.paragraphs),
        'table_count': len(doc.tables),
        'structure': content_structure
    }

if __name__ == "__main__":
    template_path = r"D:\信息工程学院2026届毕业生毕业设计 (1) (2).docx"
    
    try:
        result = read_template(template_path)
        print("\n\n[成功] 模板读取成功！")
        print(f"共发现 {result['paragraph_count']} 个段落，{result['table_count']} 个表格")
    except Exception as e:
        print(f"\n[错误] 读取失败: {e}")
        import traceback
        traceback.print_exc()
