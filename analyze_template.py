# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def analyze_template_detailed(template_path):
    """详细分析模板的格式要求"""
    doc = Document(template_path)
    
    output = []
    output.append("=" * 100)
    output.append("信息工程学院2026届毕业设计模板 - 详细格式分析报告")
    output.append("=" * 100)
    
    # 1. 页面设置
    output.append("\n\n" + "="*80)
    output.append("【一、页面布局设置】")
    output.append("="*80)
    
    if doc.sections:
        section = doc.sections[0]
        output.append(f"\n页面尺寸:")
        output.append(f"  宽度: {section.page_width.cm:.2f} cm ({section.page_width.inches:.2f} inches)")
        output.append(f"  高度: {section.page_height.cm:.2f} cm ({section.page_height.inches:.2f} inches)")
        output.append(f"\n页边距:")
        output.append(f"  上边距: {section.top_margin.cm:.2f} cm")
        output.append(f"  下边距: {section.bottom_margin.cm:.2f} cm")
        output.append(f"  左边距: {section.left_margin.cm:.2f} cm")
        output.append(f"  右边距: {section.right_margin.cm:.2f} cm")
        if section.gutter:
            output.append(f"  装订线: {section.gutter.cm:.2f} cm")
        
        # 页眉页脚
        if section.header:
            header_text = section.header.paragraphs[0].text if section.header.paragraphs else ""
            output.append(f"\n页眉内容: '{header_text}'")
        if section.footer:
            footer_text = footer.paragraphs[0].text if (footer := section.footer) and footer.paragraphs else ""
            output.append(f"页脚内容: '{footer_text}'")
    
    # 2. 提取关键文本内容
    output.append("\n\n" + "="*80)
    output.append("【二、模板核心内容摘要】")
    output.append("="*80)
    
    key_content = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 10:
            style = para.style.name if para.style else "Normal"
            
            # 获取字体详细信息
            font_details = {}
            if para.runs:
                run = para.runs[0]
                font_details['name'] = run.font.name or '继承'
                font_details['size'] = f"{run.font.size.pt}pt" if run.font.size else '继承'
                font_details['bold'] = run.font.bold
                font_details['italic'] = run.font.italic
            
            # 行距和段间距
            pf = para.paragraph_format
            line_spacing = pf.line_spacing
            space_before = pf.space_before.pt if pf.space_before else 0
            space_after = pf.space_after.pt if pf.space_after else 0
            
            entry = {
                'text': text[:200] + ('...' if len(text) > 200 else ''),
                'style': style,
                'font': font_details,
                'line_spacing': f"{line_spacing}" if line_spacing else "单倍",
                'space_before': f"{space_before}pt",
                'space_after': f"{space_after}pt"
            }
            key_content.append(entry)
    
    # 输出前50条重要内容
    output.append(f"\n共发现 {len(key_content)} 条有效段落，显示前60条:\n")
    for i, item in enumerate(key_content[:60], 1):
        output.append(f"\n--- 段落 {i} ---")
        output.append(f"样式: {item['style']}")
        output.append(f"字体: {item['font']['name']}, 字号: {item['font']['size']}, 粗体: {item['font']['bold']}")
        output.append(f"行距: {item['line_spacing']}, 段前: {item['space_before']}, 段后: {item['space_after']}")
        output.append(f"内容: {item['text']}")
    
    # 3. 表格分析
    output.append("\n\n" + "="*80)
    output.append("【三、表格结构】")
    output.append("="*80)
    
    if doc.tables:
        for t_idx, table in enumerate(doc.tables, 1):
            output.append(f"\n表格 {t_idx}:")
            output.append(f"  维度: {len(table.rows)} 行 × {len(table.columns)} 列")
            
            # 显示所有行
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ')[:30] for cell in row.cells]
                output.append(f"  第{r_idx+1}行: {' | '.join(cells)}")
    
    # 4. 样式汇总
    output.append("\n\n" + "="*80)
    output.append("【四、使用的样式列表】")
    output.append("="*80)
    
    styles_used = set()
    for para in doc.paragraphs:
        if para.style:
            styles_used.add(para.style.name)
    
    output.append("\n文档中使用的样式:")
    for style in sorted(styles_used):
        output.append(f"  - {style}")
    
    result = '\n'.join(output)
    
    # 保存到文件
    with open('template_analysis.txt', 'w', encoding='utf-8') as f:
        f.write(result)
    
    print(result)
    print("\n\n[完成] 分析结果已保存至 template_analysis.txt")
    
    return result

if __name__ == "__main__":
    template_path = r"D:\信息工程学院2026届毕业生毕业设计 (1) (2).docx"
    
    try:
        analyze_template_detailed(template_path)
    except Exception as e:
        print(f"[错误] 分析失败: {e}")
        import traceback
        traceback.print_exc()
